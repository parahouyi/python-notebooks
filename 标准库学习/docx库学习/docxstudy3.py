from docx import Document
from docx.shared import Cm

doc = Document(r"E:\2021年城建环资工作\test\333.docx")
print(doc.paragraphs)
i = 1
# for paragraph in doc.paragraphs:
#     if paragraph:
#         print(f'段落{i}:')
#         paragraph.text =paragraph.text.replace('全国人民代表大会常委会委员会', '全国人大')
#         print(paragraph.text)
#         i += 1
#
#     else:
#         print('这是无关的东西')
#
# # doc.save(r"E:\2022 年城建环资工作\全国人大环资委会议\全国人大城建环资委会议情况报告上（报省委0210杨主任修改稿）.docx")
# # doc.add_heading('this is the title',level=1)
#
# doc.add_paragraph('this is the third paragraph')
# paragraph3 = doc.add_paragraph()
# paragraph3.add_run("我被加粗了文字块儿").bold = True
# paragraph3.add_run("，我是普通文字块儿，")
# paragraph3.add_run("我是斜体文字块儿").italic = True


# doc.add_picture(r"C:\Users\Administrator\Pictures\思维导图\考试脑科学.png",width=Cm(5),height=Cm(5))
# doc.add_table(2,5)
from docx import Document

# list1 = [
#     ["姓名", "性别", "家庭地址"],
#     ["唐僧", "男", "湖北省"],
#     ["孙悟空", "男", "北京市"],
#     ["猪八戒", "男", "广东省"],
#     ["沙和尚", "男", "湖南省"]
# ]
# list2 = [
#     ["姓名", "性别", "家庭地址"],
#     ["貂蝉", "女", "河北省"],
#     ["杨贵妃", "女", "贵州省"],
#     ["西施", "女", "山东省"]
# ]
#
# table1 = doc.add_table(rows=5, cols=3)
# for row in range(5):
#     cells = table1.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list1[row][col])
# doc.add_paragraph("-----------------------------------------------------------")
# table2 = doc.add_table(rows=4, cols=3)
# for row in range(4):
#     cells = table2.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list2[row][col])

