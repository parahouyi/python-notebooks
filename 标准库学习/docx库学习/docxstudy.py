from docx import Document
from lxml import etree
from mytest.printing import printing
import re


doc = Document('E:/2022年城建环资工作/收文办理/收文015.docx')  # type:Document
# print(doc)


    # try:
    #     # print(para.text)
    #
    #     for run in para.runs:
    #         # run.replace('会议', '座谈会')
    #         # print(run)
    #         # print(run.text)
    #         # print(run.style)
# for para in doc.paragraphs:
#     print('段落：', para.text)
#     for run in para.runs:
#         print(run.text)

text = ''
for tab in doc.tables:

    rows_num = len(tab.rows)
    columns_num = len(tab.columns)
    cell_set = set()
    for i in range(rows_num):
        for j in range(columns_num):
            cell_set.add(tab.cell(i,j))
    print(cell_set)
    for i in cell_set:
        text += i.text
# print(text)


            # count += 1

            # print(tab.cell(i,j).text)




pattern = re.compile(r'<w:t>(.*)</w:t>')

# for run in doc.sections:
#     print(run)
